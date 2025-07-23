import os
from docx import Document as DocxDocument
from typing import List
from llama_index.core import Document
import logging
import os
from extraction.base import BaseExtractor
from extraction.helper import generate_metadata

logging.basicConfig(level=logging.INFO)

class ExtractDOCX(BaseExtractor):
    def extract_and_chunk(self, file_path: str) -> List[Document]:
        """Extract and chunk DOCX content with metadata."""
        if not self.validate_file(file_path):
            return []

        self.logger.info(f"Extracting and chunking: {file_path}")
        source = os.path.basename(file_path)
        ext = os.path.splitext(file_path)[-1][1:].lower()

        try:
            doc = DocxDocument(file_path)
            all_nodes = []
            current_section = None

            # Extract paragraphs and headings
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    if para.style.name.startswith("Heading"):
                        current_section = para.text.strip()
                    paragraphs.append(para.text.strip())

            self.logger.info(f"Found {len(paragraphs)} non-empty paragraphs")

            # Extract tables
            table_texts = []
            for table_idx, table in enumerate(doc.tables):
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_texts.append(row_text)
            self.logger.info(f"Extracted {len(table_texts)} table rows")

            # Combine content
            full_text = "\n".join(paragraphs + table_texts)
            if not full_text.strip():
                self.logger.warning("No content found in DOCX document.")
                return []

            # Create document
            doc = self.create_document(full_text, metadata=generate_metadata(
                source=source,
                index=0,
                max_index=1,
                file_format=ext,
                section_title=current_section
            ))
            if doc:
                nodes = self.chunk_document(doc)
                for i, node in enumerate(nodes):
                    node.metadata.update({"chunk_index": i, "total_chunks": len(nodes)})
                all_nodes.extend(nodes)

            self.logger.info(f"Extracted {len(all_nodes)} chunks from {file_path}")
            return all_nodes
        except Exception as e:
            self.logger.error(f"Error processing DOCX file: {e}")
            return []